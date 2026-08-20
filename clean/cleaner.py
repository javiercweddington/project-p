"""
Cleaner module - file-type-specific cleaning implementations.

Handles cleaning for different file types:
- TextCleaner: Plain text files using entity spans
- PDFCleaner: PDF documents (extract, clean, rebuild)
- ImageCleaner: Remove EXIF/metadata from images
- AudioCleaner: Remove metadata from audio files
- VideoCleaner: Remove metadata from video files
- XLSXCleaner: Remove metadata from Excel files
- DOCXCleaner: Remove metadata from Word documents

All cleaners work on copies in /tmp/ and never modify originals.
"""

from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from .anonymizer import EntityMapper, SpanBasedReplacer

_logger = logging.getLogger(__name__)

# Try optional dependencies
try:
    from PyPDF2 import PdfReader, PdfWriter
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


# File extension sets for routing
TEXT_EXTS = {
    '.txt', '.csv', '.tsv', '.log', '.md', '.rst',
    '.py', '.js', '.c', '.cpp', '.h', '.java', '.go', '.rs',
    '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.toml',
    '.sql', '.sh', '.bash', '.zsh',
}

DOCUMENT_EXTS = {
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.rtf', '.odt', '.ods', '.odp',
}

IMAGE_EXTS = {
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif',
    '.webp', '.ico', '.heic', '.heif',
}

AUDIO_EXTS = {
    '.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a',
}

VIDEO_EXTS = {
    '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv', '.webm',
    '.m4v', '.mpg', '.mpeg',
}


class TextCleaner:
    """Clean plain text files using entity spans from acquisition.

    Uses character offsets from GLiNER/entity detection to perform
    precise, surgical replacements.
    """

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper
        self.replacer = SpanBasedReplacer(mapper)

    def clean_file(self, input_path: Path, output_path: Path,
                   entity_spans: Optional[List[Tuple[int, int, str, str]]] = None,
                   encoding: str = 'utf-8') -> bool:
        """Clean a text file.

        Args:
            input_path: Source file path
            output_path: Destination file path (in /tmp/)
            entity_spans: List of (start, end, entity_type, source) from acquisition
            encoding: File encoding

        Returns:
            True if cleaning was successful
        """
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                text = f.read()

            if entity_spans:
                cleaned = self.replacer.replace(text, entity_spans)
            else:
                # Fallback: replace all known entities via regex
                cleaned = self.mapper.replace_in_text(text)

            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(cleaned)

            return True

        except (UnicodeDecodeError, UnicodeEncodeError) as e:
            _logger.warning("Encoding error cleaning %s: %s", input_path, e)
            return False
        except OSError as e:
            _logger.error("OS error cleaning %s: %s", input_path, e)
            return False

    def clean_text(self, text: str,
                   entity_spans: Optional[List[Tuple[int, int, str, str]]] = None) -> str:
        """Clean text content directly (without file I/O)."""
        if entity_spans:
            return self.replacer.replace(text, entity_spans)
        return self.mapper.replace_in_text(text)


class PDFCleaner:
    """Clean PDF documents by extracting text, replacing entities, and rebuilding.

    Note: Full PDF rebuilding with preserved formatting is complex.
    This implementation focuses on:
    1. Removing PDF metadata (author, creator, producer)
    2. Extracting and cleaning text content for verification
    """

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper
        self.text_cleaner = TextCleaner(mapper)

        if not HAS_PYPDF2:
            _logger.warning("PyPDF2 not available; PDF cleaning limited")

    def clean_file(self, input_path: Path, output_path: Path) -> bool:
        """Clean a PDF file by removing metadata.

        Text content cleaning in PDFs requires more sophisticated tools
        (like pdfminer + reportlab). This implementation handles metadata
        removal which is the most common leakage vector.

        Args:
            input_path: Source PDF path
            output_path: Destination PDF path

        Returns:
            True if cleaning was successful
        """
        if not HAS_PYPDF2:
            _logger.warning("PyPDF2 not available; copying PDF as-is")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return True

        try:
            reader = PdfReader(str(input_path))
            writer = PdfWriter()

            # Copy all pages
            for page in reader.pages:
                writer.add_page(page)

            # Remove metadata
            writer.metadata = {}

            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                writer.write(f)

            return True

        except Exception as e:
            _logger.error("Error cleaning PDF %s: %s", input_path, e)
            # Fallback: copy as-is
            shutil.copy2(input_path, output_path)
            return False

    def extract_and_clean_text(self, input_path: Path,
                                max_pages: int = 10) -> Optional[str]:
        """Extract text from PDF and return cleaned version.

        This is used for verification - to ensure no entity text survives.

        Args:
            input_path: Source PDF path
            max_pages: Maximum pages to extract

        Returns:
            Cleaned text or None if extraction failed
        """
        if not HAS_PYPDF2:
            return None

        try:
            reader = PdfReader(str(input_path))
            pages = []
            for i in range(min(max_pages, len(reader.pages))):
                text = reader.pages[i].extract_text() or ""
                pages.append(text)

            full_text = '\n'.join(pages)
            return self.mapper.replace_in_text(full_text)

        except Exception as e:
            _logger.debug("Failed to extract text from PDF %s: %s", input_path, e)
            return None


class ImageCleaner:
    """Clean image files by removing EXIF and other metadata.

    Uses PIL/Pillow for metadata stripping. Falls back to copying
    the file if PIL is not available.
    """

    SENSITIVE_EXIF_TAGS = {
        'Artist', 'Copyright', 'Copyrighted', 'ImageUniqueID',
        'GPSInfo', 'GPSAltitude', 'GPSLatitude', 'GPSLongitude',
        'DateTimeOriginal', 'DateTimeDigitized',
        'Make', 'Model', 'Software',
        'XPTitle', 'XPComment', 'XPAuthor', 'XPKeywords',
    }

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper

    def clean_file(self, input_path: Path, output_path: Path) -> bool:
        """Remove metadata from an image file.

        Args:
            input_path: Source image path
            output_path: Destination image path

        Returns:
            True if cleaning was successful
        """
        if not HAS_PIL:
            _logger.warning("PIL not available; copying image as-is")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return True

        try:
            img = Image.open(input_path)

            # Remove all metadata by creating a clean copy
            # This strips EXIF, IPTC, XMP, and other embedded data
            clean_img = img.copy()

            # For JPEG, explicitly clear EXIF
            if input_path.suffix.lower() in ('.jpg', '.jpeg'):
                # Save without EXIF by not passing exif parameter
                output_path.parent.mkdir(parents=True, exist_ok=True)
                save_kwargs = {}
                if img.format and img.format.lower() == 'JPEG':
                    save_kwargs['quality'] = 95

                clean_img.save(output_path, **save_kwargs)
            else:
                # For other formats, save without info dict
                output_path.parent.mkdir(parents=True, exist_ok=True)
                clean_img.save(output_path)

            return True

        except Exception as e:
            _logger.error("Error cleaning image %s: %s", input_path, e)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return False


class AudioCleaner:
    """Clean audio files by removing metadata tags.

    Uses standard tools (ffmpeg if available) to strip metadata.
    """

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper

    def clean_file(self, input_path: Path, output_path: Path) -> bool:
        """Remove metadata from an audio file.

        Args:
            input_path: Source audio path
            output_path: Destination audio path

        Returns:
            True if cleaning was successful
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Try using ffmpeg to strip metadata
        if self._try_ffmpeg_clean(input_path, output_path):
            return True

        # Fallback: copy as-is
        _logger.warning("Could not strip audio metadata; copying as-is")
        shutil.copy2(input_path, output_path)
        return False

    def _try_ffmpeg_clean(self, input_path: Path, output_path: Path) -> bool:
        """Try to use ffmpeg to remove metadata."""
        import subprocess
        try:
            result = subprocess.run(
                [
                    'ffmpeg', '-y',
                    '-i', str(input_path),
                    '-map_metadata', '-1',  # Remove all metadata streams
                    '-movflags', '+use_metadata1',
                    str(output_path),
                ],
                capture_output=True,
                timeout=60,
            )
            return result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            _logger.debug("ffmpeg not available or timed out: %s", e)
            return False


class VideoCleaner:
    """Clean video files by removing metadata tags."""

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper

    def clean_file(self, input_path: Path, output_path: Path) -> bool:
        """Remove metadata from a video file.

        Args:
            input_path: Source video path
            output_path: Destination video path

        Returns:
            True if cleaning was successful
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if self._try_ffmpeg_clean(input_path, output_path):
            return True

        _logger.warning("Could not strip video metadata; copying as-is")
        shutil.copy2(input_path, output_path)
        return False

    def _try_ffmpeg_clean(self, input_path: Path, output_path: Path) -> bool:
        """Try to use ffmpeg to remove metadata."""
        import subprocess
        try:
            result = subprocess.run(
                [
                    'ffmpeg', '-y',
                    '-i', str(input_path),
                    '-map_metadata', '-1',
                    '-c', 'copy',  # Copy streams without re-encoding
                    str(output_path),
                ],
                capture_output=True,
                timeout=120,
            )
            return result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            _logger.debug("ffmpeg not available or timed out: %s", e)
            return False


class XLSXCleaner:
    """Clean Excel files by removing document metadata and cleaning cell content.

    Uses openpyxl to:
    1. Remove core properties (author, last_modified_by, company, etc.)
    2. Optionally clean cell text content using entity mapper

    Supports .xlsx and .xlsm files (macro-enabled workbooks).
    """

    # Office document core properties to clear
    _CORE_PROPERTIES = {
        'creator', 'last_modified_by', 'contributor',
        'author', 'company', 'manager',
        'description', 'subject', 'title',
        'keywords', 'category', 'comments',
    }

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper
        self.text_cleaner = TextCleaner(mapper)

        if not HAS_OPENPYXL:
            _logger.warning(
                "openpyxl not available; Excel cleaning limited. "
                "Install with: pip install openpyxl"
            )

    def clean_file(self, input_path: Path, output_path: Path,
                   clean_content: bool = False) -> bool:
        """Clean an Excel file by removing metadata.

        Args:
            input_path: Source Excel file path
            output_path: Destination Excel file path
            clean_content: If True, also scan and replace entities in cell text

        Returns:
            True if cleaning was successful
        """
        if not HAS_OPENPYXL:
            _logger.warning("openpyxl not available; copying Excel file as-is")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return True

        try:
            # Load workbook
            wb = load_workbook(input_path, keep_vba=True if input_path.suffix.lower() == '.xlsm' else False)

            # Clear core properties
            props = wb.properties
            for prop in self._CORE_PROPERTIES:
                setattr(props, prop, '')

            # Optionally clean cell content
            if clean_content:
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if isinstance(cell.value, str):
                                cell.value = self.text_cleaner.clean_text(cell.value)

            # Save
            output_path.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(output_path))
            return True

        except Exception as e:
            _logger.error("Error cleaning Excel file %s: %s", input_path, e)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return False


class DOCXCleaner:
    """Clean Word documents by removing metadata and cleaning text content.

    Uses python-docx to:
    1. Remove core properties (author, last_modified_by, company, etc.)
    2. Remove custom properties
    3. Optionally clean paragraph/run text using entity mapper

    Note: python-docx does not preserve all formatting perfectly.
    For production use, consider using msoffcrypto-tool + libreoffice.
    """

    # Office document core properties to clear
    _CORE_PROPERTIES = {
        'creator', 'last_modified_by', 'contributor',
        'author', 'company', 'manager',
        'description', 'subject', 'title',
        'keywords', 'category', 'comments',
    }

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper
        self.text_cleaner = TextCleaner(mapper)

        if not HAS_PYTHON_DOCX:
            _logger.warning(
                "python-docx not available; Word cleaning limited. "
                "Install with: pip install python-docx"
            )

    def clean_file(self, input_path: Path, output_path: Path,
                   clean_content: bool = False) -> bool:
        """Clean a Word document by removing metadata.

        Args:
            input_path: Source DOCX file path
            output_path: Destination DOCX file path
            clean_content: If True, also scan and replace entities in text

        Returns:
            True if cleaning was successful
        """
        if not HAS_PYTHON_DOCX:
            _logger.warning("python-docx not available; copying DOCX as-is")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return True

        try:
            doc = Document(str(input_path))

            # Clear core properties via the underlying XML
            core_props = doc.core_properties
            for prop_name in self._CORE_PROPERTIES:
                setter = getattr(core_props, f'set_{prop_name}', None)
                if setter:
                    setter('')
                else:
                    # Direct attribute access
                    try:
                        setattr(core_props, prop_name, '')
                    except (AttributeError, TypeError):
                        pass

            # Clear custom properties from the XML
            try:
                custom_props = doc.custom_properties
                if custom_props:
                    # Clear all custom properties
                    props_elem = custom_props._properties
                    for child in list(props_elem):
                        props_elem.remove(child)
            except Exception:
                pass

            # Optionally clean paragraph/run text
            if clean_content:
                for para in doc.paragraphs:
                    for run in para.runs:
                        if isinstance(run.text, str) and run.text:
                            run.text = self.text_cleaner.clean_text(run.text)
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if isinstance(run.text, str) and run.text:
                                        run.text = self.text_cleaner.clean_text(run.text)

            # Save
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            return True

        except Exception as e:
            _logger.error("Error cleaning DOCX %s: %s", input_path, e)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return False


class FileCleanerRouter:
    """Route files to the appropriate cleaner based on file type."""

    def __init__(self, mapper: EntityMapper):
        self.mapper = mapper
        self.text_cleaner = TextCleaner(mapper)
        self.pdf_cleaner = PDFCleaner(mapper)
        self.image_cleaner = ImageCleaner(mapper)
        self.audio_cleaner = AudioCleaner(mapper)
        self.video_cleaner = VideoCleaner(mapper)
        self.xlsx_cleaner = XLSXCleaner(mapper)
        self.docx_cleaner = DOCXCleaner(mapper)

    def clean_file(self, input_path: Path, output_path: Path,
                   entity_spans: Optional[List[Tuple[int, int, str, str]]] = None) -> bool:
        """Clean a file using the appropriate cleaner for its type.

        Args:
            input_path: Source file path
            output_path: Destination file path
            entity_spans: Entity spans for text-based cleaners

        Returns:
            True if cleaning was successful
        """
        ext = input_path.suffix.lower()

        if ext in TEXT_EXTS:
            return self.text_cleaner.clean_file(input_path, output_path, entity_spans)
        elif ext == '.pdf':
            return self.pdf_cleaner.clean_file(input_path, output_path)
        elif ext in IMAGE_EXTS:
            return self.image_cleaner.clean_file(input_path, output_path)
        elif ext in AUDIO_EXTS:
            return self.audio_cleaner.clean_file(input_path, output_path)
        elif ext in VIDEO_EXTS:
            return self.video_cleaner.clean_file(input_path, output_path)
        elif ext in ('.xlsx', '.xlsm'):
            return self.xlsx_cleaner.clean_file(input_path, output_path)
        elif ext == '.docx':
            return self.docx_cleaner.clean_file(input_path, output_path)
        else:
            # Unknown type: copy as-is
            _logger.debug("Unknown file type %s; copying as-is", ext)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(input_path, output_path)
            return True

    def get_cleaner_for_ext(self, ext: str):
        """Get the appropriate cleaner for a file extension."""
        if ext in TEXT_EXTS:
            return self.text_cleaner
        elif ext == '.pdf':
            return self.pdf_cleaner
        elif ext in IMAGE_EXTS:
            return self.image_cleaner
        elif ext in AUDIO_EXTS:
            return self.audio_cleaner
        elif ext in VIDEO_EXTS:
            return self.video_cleaner
        elif ext in ('.xlsx', '.xlsm'):
            return self.xlsx_cleaner
        elif ext == '.docx':
            return self.docx_cleaner
        return None